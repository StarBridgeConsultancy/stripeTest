# ─── Flask Core ─────────────────────────────────────────
from flask import (
    Flask, render_template, request, jsonify, redirect,
    url_for, session, flash, send_file
)

# ─── Flask Extensions ──────────────────────────────────
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from flask_login import (
    LoginManager, login_required, current_user
)
from flask_mail import Mail, Message
from flask_wtf import FlaskForm

# ─── WTForms ────────────────────────────────────────────
from wtforms import StringField, TextAreaField, SelectField, SubmitField
from wtforms.validators import DataRequired, Email, URL, Optional

# ─── Utilities & Standard Library ───────────────────────
from werkzeug.utils import secure_filename
from datetime import datetime, timezone, timedelta
from io import BytesIO
import os, re
from flask import request, session, jsonify
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ─── External Services ──────────────────────────────────
from fpdf import FPDF
from docx import Document
import stripe
import google.generativeai as genai

# ─── App Initialization ─────────────────────────────────
app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Replace securely in production

# ─── Configurations ─────────────────────────────────────
app.config.update(
    SQLALCHEMY_DATABASE_URI='postgresql://myapp_r6n3_user:H6R5XiPx4EL8gevrbW4ySEjZIhf72U3y@dpg-d1i13dodl3ps73b1ktc0-a.oregon-postgres.render.com/myapp_r6n3',
    SQLALCHEMY_TRACK_MODIFICATIONS=False,
    MAIL_DEFAULT_SENDER='starbridgeconsultancy@gmail.com',
    MAIL_SERVER='smtp.gmail.com',
    MAIL_PORT=587,
    MAIL_USE_TLS=True,
    MAIL_USERNAME='starbridgeconsultancy@gmail.com',
    MAIL_PASSWORD='gooh wfay uxur yhsa',
    UPLOAD_FOLDER='uploads'
)
ALLOWED_EXTENSIONS = {"pdf", "txt"}
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

# ─── Service Initialization ─────────────────────────────
db = SQLAlchemy(app)
migrate = Migrate(app, db)
mail = Mail(app)

stripe.api_key = os.environ.get('STRIPE_SECRET_KEY')
STRIPE_PUBLISHABLE_KEY = os.environ.get('STRIPE_PUBLISHABLE_KEY')

# ─── Login Manager Setup ────────────────────────────────
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))

@login_manager.unauthorized_handler
def unauthorized_callback():
    if request.is_json or request.path.startswith("/chatbot-api"):
        return jsonify({"error": "Unauthorized", "login_required": True}), 401
    return redirect(url_for("login"))

# ─── Gemini Configuration ───────────────────────────────
API_KEY = "AIzaSyA_84rmTgnFdvzjpFdB8p3xYoziCVbcEic"  # ⚠️ Move to .env in production
genai.configure(api_key=API_KEY)
model = genai.GenerativeModel("gemini-2.0-flash")
chat = model.start_chat()

# ─── Models ─────────────────────────────────────────────

class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(100), unique=True, nullable=False)
    password = db.Column(db.String(100), nullable=False)
    is_subscribed = db.Column(db.Boolean, default=False)

    # Profile Info
    full_name = db.Column(db.String(100))
    phone = db.Column(db.String(20))
    address = db.Column(db.String(200))
    country = db.Column(db.String(100))
    city = db.Column(db.String(100))
    preferred_job_type = db.Column(db.String(50))  # Remote / Hybrid / Onsite
    preferred_industries = db.Column(db.Text)      # Eg: IT, Healthcare, Marketing
    skills = db.Column(db.Text)
    experience = db.Column(db.Text)
    education = db.Column(db.Text)

    applications = db.relationship(
        'JobApplication',
        back_populates='user',
        lazy=True,
        cascade='all, delete-orphan'
    )


class Job(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(255), nullable=False)
    company = db.Column(db.String(255), nullable=False)
    location = db.Column(db.String(255))
    country = db.Column(db.String(100))
    industry = db.Column(db.String(255))
    location_type = db.Column(db.String(50))
    link = db.Column(db.String(500), nullable=False)
    email = db.Column(db.String(255))  # Contact email
    description = db.Column(db.Text)
    posted_date = db.Column(db.DateTime, default=datetime.utcnow)
    expiry_days = db.Column(db.Integer, default=30)

    applications = db.relationship(
        'JobApplication',
        back_populates='job',
        lazy=True,
        cascade='all, delete-orphan'
    )


class JobApplication(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id', ondelete='CASCADE'), nullable=False)
    job_id = db.Column(db.Integer, db.ForeignKey('job.id', ondelete='CASCADE'), nullable=False)
    applied_date = db.Column(db.DateTime, default=datetime.utcnow)

    user = db.relationship('User', back_populates='applications')
    job = db.relationship('Job', back_populates='applications')


class Course(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    academy = db.Column(db.String(100), nullable=False)
    fee = db.Column(db.Float, nullable=False)
    start_date = db.Column(db.Date, nullable=False)
    end_date = db.Column(db.Date, nullable=False)
    description = db.Column(db.Text)
    link = db.Column(db.String(255))  # New field for course URL

    registrations = db.relationship('CourseRegistration', back_populates='course', cascade='all, delete-orphan')



class CourseRegistration(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id', ondelete='CASCADE'))
    course_id = db.Column(db.Integer, db.ForeignKey('course.id', ondelete='CASCADE'))
    registered_at = db.Column(db.DateTime, default=datetime.utcnow)

    user = db.relationship('User', backref='course_registrations')
    course = db.relationship('Course', back_populates='registrations')



# ─── WTForms ────────────────────────────────────────────

class JobForm(FlaskForm):
    title = StringField('Title', validators=[DataRequired()])
    company = StringField('Company', validators=[DataRequired()])
    location = StringField('Location', validators=[Optional()])
    email = StringField('Application Email', validators=[Optional(), Email()])
    
    country = SelectField(
        'Country',
        choices=[
            ('', 'Select Country'),
            ('United Kingdom', 'United Kingdom'),
            ('United States', 'United States'),
            ('Canada', 'Canada'),
            ('India', 'India'),
            ('Germany', 'Germany'),
            ('Australia', 'Australia'),
            # Add more if needed
        ],
        validators=[Optional()]
    )
    
    industry = StringField('Industry', validators=[Optional()])
    location_type = SelectField(
        'Location Type',
        choices=[
            ('Remote', 'Remote'),
            ('Onsite', 'Onsite'),
            ('Hybrid', 'Hybrid')
        ]
    )
    link = StringField('Job Link', validators=[DataRequired(), URL()])
    description = TextAreaField('Description', validators=[Optional()])
    submit = SubmitField('Add Job')


class CourseForm(FlaskForm):
    title = StringField('Course Title', validators=[DataRequired()])
    academy = StringField('Academy Offering Course', validators=[DataRequired()])
    fee = StringField('Course Fee (e.g., $99)', validators=[DataRequired()])
    start_date = StringField('Start Date (YYYY-MM-DD)', validators=[DataRequired()])
    end_date = StringField('End Date (YYYY-MM-DD)', validators=[DataRequired()])
    link = StringField('Course Link (URL)')  # New form field
    description = TextAreaField('Course Description')
    submit = SubmitField('Add Course')

    

# ─── PDF CV Generator Class ─────────────────────────────

class CVPDF(FPDF):
    def header(self):
        self.set_font("Helvetica", "B", 18)
        self.set_text_color(0, 51, 102)
        self.cell(0, 10, self.name, ln=True)

        self.set_font("Helvetica", "", 10)
        self.set_text_color(0)
        if self.location:
            self.cell(0, 6, self.location, ln=True)
        if self.phone:
            self.cell(0, 6, f"Phone: {self.phone}", ln=True)
        if self.email:
            self.cell(0, 6, f"Email: {self.email}", ln=True, link=f"mailto:{self.email}")
        if self.linkedin:
            self.set_text_color(0, 102, 204)
            self.cell(0, 6, f"LinkedIn: {self.linkedin}", ln=True, link=self.linkedin)
        self.set_text_color(0)
        self.ln(5)

    def section_title(self, title):
        self.set_fill_color(224, 224, 224)
        self.set_text_color(0, 51, 102)
        self.set_font("Helvetica", "B", 11)
        self.cell(0, 8, f"  {title}", ln=True, fill=True)
        self.set_text_color(0)
        self.set_font("Helvetica", "", 10)

    def section_body(self, text):
        if text:
            self.multi_cell(0, 6, text.strip())
            self.ln(2)

# ─── Email Body Generator ───────────────────────────────

def generate_email_body(user, job):
    return f"""
Dear {job.company} Hiring Team,

I hope this message finds you well.

My name is {user.full_name}, and I am writing to express my strong interest in the position of {job.title} at your esteemed organization. I believe my skills and experience make me a great fit for this opportunity.

Please find my CV attached for your consideration. I look forward to the opportunity to contribute to your team.

Sincerely,  
{user.full_name}  
{user.email}
"""

# ─── Job Cleanup Utility ────────────────────────────────

def cleanup_expired_jobs():
    now = datetime.now(timezone.utc)
    expired_jobs = []

    for job in Job.query.all():
        if job.posted_date is None:
            expired_jobs.append(job)
            continue

        expiry = job.expiry_days if job.expiry_days is not None else 30
        posted_date = job.posted_date

        if posted_date.tzinfo is None:
            posted_date = posted_date.replace(tzinfo=timezone.utc)

        if posted_date + timedelta(days=expiry) < now:
            expired_jobs.append(job)

    for job in expired_jobs:
        db.session.delete(job)
    db.session.commit()

# ─── Notification Utility ───────────────────────────────

def notify_users_about_new_job(job):
    users = User.query.all()
    subject = f"New Job Posted: {job.title} at {job.company}"

    for user in users:
        if user.email:
            msg = Message(subject, recipients=[user.email])
            msg.body = f"""
Hi {user.full_name},

A new job has just been posted on the platform:

Title: {job.title}
Company: {job.company}
Location: {job.location}, {job.country}
Industry: {job.industry}
Type: {job.location_type}
Link: {job.link}

Description:
{job.description}

Log in to your dashboard to apply or view more jobs.

Best regards,  
GoGetJobs Team
"""
            try:
                mail.send(msg)
            except Exception as e:
                print(f"Error sending to {user.email}: {e}")
# ─── Routes: Home & Auth ────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        existing_user = User.query.filter_by(email=email).first()
        if existing_user:
            return "User already exists."
        new_user = User(email=email, password=password)
        db.session.add(new_user)
        db.session.commit()
        return redirect('/login')
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        user = User.query.filter_by(email=email, password=password).first()
        if user:
            session['user_id'] = user.id
            return redirect('/jobs')
        return "Invalid credentials"
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    return redirect('/')

# ─── Routes: Dashboard & Profile ────────────────────────

@app.route('/dashboard')
def dashboard():
    if 'user_id' not in session:
        return redirect('/login')
    
    user = db.session.get(User, session['user_id'])
    if not user:
        flash("User not found.", "error")
        return redirect('/login')
    
    # Load relationships
    applications = user.applications  # Job applications
    registered_courses = user.course_registrations  # Courses registered
    
    return render_template(
        'dashboard.html',
        user=user,
        applications=applications,
        registered_courses=registered_courses
    )


@app.route('/profile', methods=['GET', 'POST'])
def profile():
    if 'user_id' not in session:
        return redirect('/login')
    user = db.session.get(User, session['user_id'])
    if request.method == 'POST':
        user.full_name = request.form.get('full_name')
        user.phone = request.form.get('phone')
        user.address = request.form.get('address')
        user.country = request.form.get('country')
        user.city = request.form.get('city')
        user.preferred_job_type = request.form.get('preferred_job_type')
        user.preferred_industries = ','.join(request.form.getlist('preferred_industries'))
        user.skills = request.form.get('skills')
        user.experience = request.form.get('experience')
        user.education = request.form.get('education')
        db.session.commit()
        return redirect('/dashboard')
    return render_template('profile.html', user=user)

# ─── Routes: Job Management ─────────────────────────────

@app.route('/admin/add-job', methods=['GET', 'POST'])
def add_job():
    form = JobForm()
    if form.validate_on_submit():
        job = Job(
            title=form.title.data,
            company=form.company.data,
            location=form.location.data,
            country=form.country.data,
            industry=form.industry.data,
            location_type=form.location_type.data,
            link=form.link.data,
            email=form.email.data,
            description=form.description.data
        )
        db.session.add(job)
        db.session.commit()
        notify_users_about_new_job(job)
        flash('Job added and users notified!', 'success')
        return redirect(url_for('add_job'))
    return render_template('admin.html', form=form)

@app.route('/admin/delete-job/<int:job_id>', methods=['POST'])
def delete_job(job_id):
    job = Job.query.get_or_404(job_id)
    db.session.delete(job)
    db.session.commit()
    flash('Job deleted successfully', 'success')
    return redirect(url_for('add_job'))

# ─── Routes: Jobs Page & Application ────────────────────

@app.route('/jobs', methods=['GET', 'POST'])
def jobs():
    if 'user_id' not in session:
        return redirect('/login')
    user = db.session.get(User, session['user_id'])
    if not user:
        flash("User not found.", "error")
        return redirect('/login')

    applied_job_ids = [app.job_id for app in user.applications]

    user_skills = set(map(str.strip, user.skills.lower().split(','))) if user.skills else set()
    user_industries = set(map(str.strip, user.preferred_industries.lower().split(','))) if user.preferred_industries else set()
    user_country = user.country.lower() if user.country else ""

    available_jobs = Job.query.filter(~Job.id.in_(applied_job_ids)).all()

    def match_job(job):
        job_industry = (job.industry or '').strip().lower()
        job_location_type = (job.location_type or '').strip().lower()
        job_country = (job.country or '').strip().lower()

        if job_industry not in user_industries:
            return False
        if job_location_type in ['remote', 'international']:
            return True
        if job_country == user_country:
            return True
        return False

    filtered_jobs = [job for job in available_jobs if match_job(job)]
    applied_jobs = Job.query.filter(Job.id.in_(applied_job_ids)).all()

    return render_template(
        'jobs.html',
        user=user,
        applied_jobs=applied_jobs,
        available_jobs=filtered_jobs,
        is_subscribed=user.is_subscribed
    )

@app.route('/apply', methods=['POST'])
def apply():
    if 'user_id' not in session:
        return redirect('/login')
    try:
        job_id = int(request.form['job_id'])
    except (ValueError, KeyError):
        flash("Invalid job selection.", "error")
        return redirect('/jobs')

    user = db.session.get(User, session['user_id'])
    job = Job.query.get(job_id)

    if not job:
        flash("Job does not exist.", "error")
        return redirect('/jobs')

    existing_application = JobApplication.query.filter_by(user_id=user.id, job_id=job.id).first()
    if existing_application:
        flash("You have already applied to this job.", "info")
        return redirect('/jobs')

    new_application = JobApplication(user_id=user.id, job_id=job.id)
    db.session.add(new_application)
    db.session.commit()
    flash("Applied successfully!", "success")

    cv_pdf = generate_cv(user)
    email_body = generate_email_body(user, job)

    try:
        msg = Message(
            subject=f"Job Application for {job.title}",
            recipients=[job.email],
            body=email_body,
            reply_to=user.email
        )
        msg.attach("CV.pdf", "application/pdf", cv_pdf.getvalue())
        mail.send(msg)
        flash("Email sent to the employer!", "success")
    except Exception as e:
        print("Email sending failed:", e)
        flash("Failed to send application email.", "error")

    return redirect('/jobs')

# ─── Routes: CV & Styling ───────────────────────────────

@app.route('/styled-cv')
def styled_cv():
    if 'user_id' not in session:
        return redirect('/login')
    user = db.session.get(User, session['user_id'])
    if not user:
        flash("User not found", "error")
        return redirect('/dashboard')

    pdf = CVPDF()
    pdf.name = user.full_name or ""
    pdf.location = f"{user.city}, {user.country}" if user.city and user.country else ""
    pdf.phone = user.phone or ""
    pdf.email = user.email or ""
    pdf.linkedin = ""  # optional
    pdf.add_page()

    if user.experience:
        pdf.section_title("Professional Summary")
        pdf.section_body(user.experience)
    if user.skills:
        pdf.section_title("Skills")
        pdf.section_body(user.skills)
    if user.preferred_industries:
        pdf.section_title("Preferred Industries")
        pdf.section_body(user.preferred_industries)
    if user.preferred_job_type:
        pdf.section_title("Preferred Job Type")
        pdf.section_body(user.preferred_job_type)
    if user.education:
        pdf.section_title("Education")
        pdf.section_body(user.education)

    buffer = BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return send_file(
        buffer,
        mimetype="application/pdf",
        as_attachment=True,
        download_name=f"{user.full_name or 'cv'}_CV.pdf"
    )

# ─── Routes: Subscription & Payment ─────────────────────

@app.route('/subscribe')
def subscribe():
    if 'user_id' not in session:
        return redirect('/login')
    user = db.session.get(User, session['user_id'])
    checkout_session = stripe.checkout.Session.create(
        customer_email=user.email,
        payment_method_types=['card'],
        line_items=[{
            'price_data': {
                'currency': 'usd',
                'unit_amount': 1000,
                'product_data': {'name': 'Monthly Access'}
            },
            'quantity': 1
        }],
        mode='payment',
        success_url=url_for('payment_success', _external=True),
        cancel_url=url_for('payment_cancel', _external=True)
    )
    return redirect(checkout_session.url, code=303)

@app.route('/payment-success')
def payment_success():
    if 'user_id' in session:
        user = db.session.get(User, session['user_id'])
        user.is_subscribed = True
        db.session.commit()
    return redirect('/jobs')

@app.route('/payment-cancel')
def payment_cancel():
    return "Payment was cancelled. <a href='/'>Go back</a>"

# ─── Routes: Chatbot Interface ──────────────────────────

@app.route('/chatbot-ui')
def chatbot_ui():
    return render_template('chatbot.html')
# ─── Global CV Buffer (In-Memory) ───────────────────────
latest_cv = BytesIO()

# ─── Chatbot API Endpoint ───────────────────────────────


@app.route('/chatbot-api', methods=['POST'])
def chatbot_api():
    global latest_cv

    if 'user_id' not in session:
        return jsonify({"reply": "❌ Please log in to use the chatbot."}), 401

    user = db.session.get(User, session['user_id'])
    if not user:
        return jsonify({"reply": "❌ User not found. Please log in again."}), 401

    data = request.get_json()
    user_message = data.get("message", "").strip()

    if not user_message:
        return jsonify({
            "reply": "Please enter a job-related question or paste a job description."
        })

    # Extract user profile
    full_name = user.full_name or "Anonymous"
    phone = user.phone or "N/A"
    email = user.email or "N/A"
    city = user.city or "N/A"
    country = user.country or "N/A"
    preferred_job_type = user.preferred_job_type or "N/A"
    preferred_industries = user.preferred_industries or "N/A"
    skills = user.skills or "N/A"
    experience = user.experience or "N/A"
    education = user.education or "N/A"

    # Determine if message is a job description
    is_job_description = any(
        keyword in user_message.lower()
        for keyword in ['requirements', 'job description', 'we are looking for', 'qualifications']
    ) or len(user_message) > 300

    if is_job_description:
        prompt = f"""
Generate a professional CV in plain text. Tailor it to the job description below:

Job Description:
{user_message}

User Profile:
Name: {full_name}
Phone: {phone}
Email: {email}
Location: {city}, {country}
Preferred Job Type: {preferred_job_type}
Preferred Industries: {preferred_industries}
Skills: {skills}
Experience: {experience}
Education: {education}

Format:
- Start with contact info
- Then PROFESSIONAL SUMMARY
- Then EXPERIENCE
- Then SKILLS
- Then EDUCATION
- Use plain text and dashes (-) for bullets.
"""

        try:
            response = model.generate_content(prompt)
            text_cv = response.text.strip()
            lines = text_cv.split('\n')

            # Start creating styled document
            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)

            # ─── Helper Functions ────────────────────────
            def add_heading(doc, text):
                heading = doc.add_paragraph()
                run = heading.add_run(text)
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(44, 43, 100)  # Deep Blue
                run.font.name = 'Calibri'
                heading.space_after = Pt(4)

            def add_bullet_points(doc, text):
                for line in text.split('\n'):
                    if line.strip().startswith('-'):
                        doc.add_paragraph(line.strip(), style='List Bullet')

            def add_horizontal_line(paragraph):
                p = paragraph._p
                pPr = p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), 'C1272D')  # Red line
                pBdr.append(bottom)
                pPr.append(pBdr)

            # ─── Build Document ──────────────────────────

            # Header
            doc.add_paragraph(lines[0].strip()).runs[0].bold = True
            doc.add_paragraph(lines[1].strip())

            add_horizontal_line(doc.add_paragraph())  # Line separator

            # Parse rest of the CV
            current_section = None
            buffer = []

            for line in lines[2:]:
                if line.strip().isupper():  # New section
                    if current_section and buffer:
                        add_bullet_points(doc, '\n'.join(buffer))
                        buffer = []
                    add_heading(doc, line.strip())
                    current_section = line.strip()
                else:
                    buffer.append(line)

            if current_section and buffer:
                add_bullet_points(doc, '\n'.join(buffer))

            # Save document
            latest_cv = BytesIO()
            doc.save(latest_cv)
            latest_cv.seek(0)

            return jsonify({
                "reply": "✅ Your tailored CV is ready. [Click here to download it](/chatbot-download-cv)"
            })

        except Exception as e:
            return jsonify({
                "reply": f"❌ Failed to generate your CV: {str(e)}"
            }), 500

    # ─── Job-Related Question ──────────────────────────

    prompt = f"""
You're a helpful job assistant bot. Answer only job-related questions.

User Profile:
Name: {full_name}
Skills: {skills}
Experience: {experience}
Education: {education}

Question: {user_message}
"""
    try:
        response = model.generate_content(prompt)
        return jsonify({"reply": response.text.strip()})
    except Exception as e:
        return jsonify({"reply": f"Error: {str(e)}"}), 500

# ─── Route: Download CV from Chatbot ────────────────────

@app.route('/chatbot-download-cv')
def chatbot_download_cv():
    if 'user_id' not in session:
        return redirect('/login')

    if not latest_cv or not latest_cv.getbuffer().nbytes:
        return "No CV generated yet.", 404

    latest_cv.seek(0)
    return send_file(
        latest_cv,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='Tailored_CV.docx'
    )
@app.route('/generate-cv')
def generate_cv():
    if 'user_id' not in session:
        return redirect('/login')

    user = db.session.get(User, session['user_id'])
    if not user:
        flash("User not found", "error")
        return redirect('/dashboard')

    prompt = f"""
Create a clean, modern, ATS-friendly CV in plain text only — no markdown, no asterisks, and no LaTeX. 
Use uppercase section headers (e.g., PROFESSIONAL SUMMARY), simple bullet points (like hyphens), and consistent spacing.

Full Name: {user.full_name}
Phone: {user.phone}
Email: {user.email}
Address: {user.address}, {user.city}, {user.country}

Professional Summary:
Write a concise 2–3 sentence summary that highlights the user's experience, goals, and key strengths.

Career Preferences:
Job Type: {user.preferred_job_type}
Industries: {user.preferred_industries}

Experience:
{user.experience}

Skills:
{user.skills}

Education:
{user.education}
"""

    try:
        response = model.generate_content(prompt)
        cv_text = response.text.strip()

        personal_lines = [
            user.full_name or "",
            user.phone or "",
            user.email or "",
            f"{user.address}, {user.city}, {user.country}"
        ]
        for line in personal_lines:
            cv_text = re.sub(re.escape(line), '', cv_text, flags=re.IGNORECASE)
        cv_text = re.sub(r'\*+|\#+', '', cv_text)

        class PDF(CVPDF): pass

        pdf = PDF()
        pdf.name = user.full_name or "CV"
        pdf.location = f"{user.city}, {user.country}" if user.city and user.country else ""
        pdf.phone = user.phone or ""
        pdf.email = user.email or ""
        pdf.linkedin = ""
        pdf.add_page()

        section_titles = re.findall(r'([A-Z][A-Z ]+):', cv_text)
        sections = re.split(r'[A-Z][A-Z ]+:\s*\n?', cv_text)

        if section_titles and len(sections) > 1:
            for title, body in zip(section_titles, sections[1:]):
                pdf.section_title(title.strip())
                pdf.section_body(body.strip())
        else:
            pdf.section_body(cv_text)

        pdf_bytes = pdf.output(dest='S')
        buffer = BytesIO(pdf_bytes)

        return send_file(
            buffer,
            mimetype="application/pdf",
            as_attachment=True,
            download_name=f"{user.full_name or 'CV'}_Gemini_Styled.pdf"
        )

    except Exception as e:
        return f"Error generating CV with Gemini: {e}"
    
    
# ─── courses ─────────────────────────────────────────
@app.route('/admin/add-course', methods=['GET', 'POST'])
def add_course():
    form = CourseForm()
    if form.validate_on_submit():
        course = Course(
            title=form.title.data,
            academy=form.academy.data,
            fee=float(form.fee.data.replace('$', '').replace(',', '')),
            start_date=datetime.strptime(form.start_date.data, "%Y-%m-%d").date(),
            end_date=datetime.strptime(form.end_date.data, "%Y-%m-%d").date(),
            description=form.description.data,
            link=form.link.data
        )
        db.session.add(course)
        db.session.commit()
        flash("Course added successfully!", "success")
        return redirect(url_for('add_course'))

    # Get all courses for display
    courses = Course.query.order_by(Course.start_date.desc()).all()
    return render_template('add_course.html', form=form, courses=courses)


# ─── Delete Course ─────────────────────────────────────────
@app.route('/admin/delete-course/<int:course_id>', methods=['POST'])
def delete_course(course_id):
    course = Course.query.get_or_404(course_id)
    db.session.delete(course)
    db.session.commit()
    flash(f"Course '{course.title}' has been deleted.", "success")
    return redirect(url_for('add_course'))  # Or wherever you're listing courses

@app.route('/admin/edit-course/<int:course_id>', methods=['GET', 'POST'])
def edit_course(course_id):
    course = Course.query.get_or_404(course_id)
    form = CourseForm(obj=course)

    if form.validate_on_submit():
        course.title = form.title.data
        course.academy = form.academy.data
        course.fee = float(form.fee.data.replace('$', '').replace(',', ''))
        course.start_date = datetime.strptime(form.start_date.data, "%Y-%m-%d").date()
        course.end_date = datetime.strptime(form.end_date.data, "%Y-%m-%d").date()
        course.description = form.description.data
        course.link = form.link.data
        db.session.commit()
        flash("Course updated successfully!", "success")
        return redirect(url_for('add_course'))

    # On GET or form errors, render dedicated edit page
    return render_template('edit_course.html', form=form, course=course)


@app.route('/courses')
def view_courses():
    if 'user_id' not in session:
        return redirect('/login')
    courses = Course.query.all()
    user = db.session.get(User, session['user_id'])
    registered_ids = {r.course_id for r in user.course_registrations}
    return render_template('courses.html', courses=courses, registered_ids=registered_ids)
@app.route('/register-course/<int:course_id>', methods=['POST'])
def register_course(course_id):
    if 'user_id' not in session:
        return redirect('/login')

    user = db.session.get(User, session['user_id'])
    if not user:
        flash("User not found", "error")
        return redirect('/courses')

    course = Course.query.get(course_id)
    if not course:
        flash("Course not found.", "error")
        return redirect('/courses')

    # Check if already registered
    existing = CourseRegistration.query.filter_by(user_id=user.id, course_id=course.id).first()
    if existing:
        flash("You are already registered for this course.", "info")
        return redirect('/courses')

    # Register user
    registration = CourseRegistration(user_id=user.id, course_id=course.id)
    db.session.add(registration)
    db.session.commit()

    # Send confirmation email
    try:
        msg = Message(
            subject=f"Registration Confirmed: {course.title}",
            recipients=[user.email],
            body=f"""
Hi {user.full_name},

You have successfully registered for the course "{course.title}" offered by {course.academy}.

📅 Start Date: {course.start_date}
📅 End Date: {course.end_date}
💰 Fee: ${course.fee:.2f}

You can access the course here:
{course.link or "No link provided"}

Thanks for registering!

— GoGetJobs Team
            """
        )
        mail.send(msg)
        flash("✅ Registered successfully. Confirmation email sent.", "success")
    except Exception as e:
        print(f"Email error: {e}")
        flash("Registered, but failed to send confirmation email.", "warning")

    return redirect('/courses')







@app.route('/admin/course-registrations')
def view_course_registrations():
    courses = Course.query.all()
    return render_template('admin_course_registrations.html', courses=courses)

# ─── App Runner ─────────────────────────────────────────


if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        cleanup_expired_jobs()
    app.run(debug=True)
